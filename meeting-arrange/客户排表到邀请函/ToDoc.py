# -*- coding: utf-8 -*-
"""
Created on Fri May 17 09:41:12 2019

@author: Xinchun

将提取结果中的 时间 上市公司 房间号 三列信息 自动写入到word
"""
import docx
import pandas as pd

#销售联系方式
agentData = pd.read_excel('销售联系方式.xlsx')

# 读入提取结果
extract = pd.read_excel('客户-提取结果.xlsx')
pidList = []
for idx in extract.index:
    ins = extract.loc[idx]['机构']
    name = extract.loc[idx]['姓名']
    pidList.append(ins+'\t'+name)
extract['个人代码'] = pidList

pidData = extract[extract['时间'] == '时间']
for pid in pidData['个人代码']:
    print(pid)
    info = extract[extract['个人代码'] == pid]
    baseDoc = docx.Document('测试底稿.docx')
    # 根据 销售、机构、客户 更替信息和生成文件名
    agentName = info.iloc[0]['对口销售']
    agentInfo = agentData[agentData['姓名'] == agentName]
    phoneNum = str(agentInfo.iloc[0].values[1])
    
    para = baseDoc.paragraphs[-5].clear()
    runTxt = para.add_run(agentName + '\t' + phoneNum)
    font = runTxt.font
    font.size = 160000
    font.bold = True
    
    para = baseDoc.paragraphs[-3].clear()
    runTxt = para.add_run(pid)
    font = runTxt.font
    font.size = 160000
    font.bold = True
    
#    baseDoc.paragraphs[-5].text = agentName + '\t' + phoneNum
#    baseDoc.paragraphs[-3].text = pid
    
    fileName = '排表结果\\'+ agentName + ' ' + pid.replace('\t', ' ') 
    # 有用的信息
    infoValid = info.loc[:,['时间','上市公司','房间号']]
    infoValid.fillna(value='',inplace=True)
    rowQnt = len(infoValid)
    colQnt = 3
#    table = baseDoc.add_table(rows=rowQnt, cols=colQnt, style='Table Grid')
    table0 = baseDoc.tables[0]
    for row in range(1, rowQnt):
        table0.add_row()
        for col in range(0, colQnt):
            table0.cell(row,col).text = infoValid.iloc[row,col]

    table0.autofit=True
    baseDoc.save(fileName +'.docx')
    
