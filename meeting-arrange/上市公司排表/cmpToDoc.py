# -*- coding: utf-8 -*-
"""
Created on Fri May 17 09:41:12 2019

@author: Xinchun

将提取结果中的 时间 上市公司 房间号 三列信息 自动写入到word
"""
import docx
import pandas as pd

# 读入提取结果
extract = pd.read_excel('上市公司提取结果.xlsx')

pidData = extract[extract['时间'] == '时间']
for pid in pidData['上市公司']:
    print(pid)
    info = extract[extract['上市公司'] == pid]
    baseDoc = docx.Document('测试底稿.docx')
    
#    baseDoc.paragraphs[-5].text = agentName + '\t' + phoneNum
#    baseDoc.paragraphs[-3].text = pid
    
    fileName = '上市公司排表结果\\'+ pid
    # 有用的信息
    infoValid = info.loc[:,['时间','房间号','客户清单']]
    infoValid.fillna(value='',inplace=True)
    rowQnt = len(infoValid)
    colQnt = 3
#    table = baseDoc.add_table(rows=rowQnt, cols=colQnt, style='Table Grid')
    table0 = baseDoc.tables[0]
    for row in range(1, rowQnt):
        table0.add_row()
        for col in range(0, colQnt):
            table0.cell(row,col).text = infoValid.iloc[row,col]

    table0.autofit=True
    baseDoc.save(fileName +'.docx')
    
